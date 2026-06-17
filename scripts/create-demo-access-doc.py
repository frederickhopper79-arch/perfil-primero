from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("ACCESOS_DEMO_PERFIL_PRIMERO.docx")

credentials = [
    ("Admin", "Administrador", "Admin", "1234", "/consola-admin"),
    ("Postulante", "Valentina Rojas", "valentina.rojas@demo.perfilprimero.cl", "Demo1234!", "/postulante"),
    ("Postulante", "Matias Araya", "matias.araya@demo.perfilprimero.cl", "Demo1234!", "/postulante"),
    ("Postulante", "Camila Fuentes", "camila.fuentes@demo.perfilprimero.cl", "Demo1234!", "/postulante"),
    ("Postulante", "Diego Morales", "diego.morales@demo.perfilprimero.cl", "Demo1234!", "/postulante"),
    ("Postulante", "Fernanda Silva", "fernanda.silva@demo.perfilprimero.cl", "Demo1234!", "/postulante"),
    ("Empresa", "Norte Digital SpA", "reclutamiento@nortedigital.demo", "Demo1234!", "/empresa"),
    ("Empresa", "Andes Operaciones SpA", "personas@andesoperaciones.demo", "Demo1234!", "/empresa"),
]

postulants = [
    ("Valentina Rojas", "Marketing digital y performance", "Google Ads, Meta Ads, GA4", "$1.200.000 - $1.800.000"),
    ("Matias Araya", "Supervisor logistico", "WMS, inventario, KPI", "$1.000.000 - $1.450.000"),
    ("Camila Fuentes", "Analista contable", "SII, ERP, Excel", "$850.000 - $1.150.000"),
    ("Diego Morales", "Desarrollador frontend", "React, Next.js, TypeScript", "$1.500.000 - $2.200.000"),
    ("Fernanda Silva", "Atencion cliente y postventa", "CRM, Zendesk, ventas", "$750.000 - $980.000"),
]

offers = [
    ("Norte Digital SpA", "Especialista SEO", "Marketing y Publicidad", "1"),
    ("Norte Digital SpA", "Analista paid media", "Marketing y Publicidad", "2"),
    ("Norte Digital SpA", "Disenador UX/UI", "Diseno", "1"),
    ("Norte Digital SpA", "Desarrollador frontend", "Tecnologia", "1"),
    ("Norte Digital SpA", "Ejecutivo comercial B2B", "Comercial", "3"),
    ("Andes Operaciones SpA", "Supervisor logistico", "Logistica", "2"),
    ("Andes Operaciones SpA", "Analista inventario", "Produccion", "1"),
    ("Andes Operaciones SpA", "Coordinador transporte", "Logistica", "1"),
    ("Andes Operaciones SpA", "Analista contable operaciones", "Administracion y Finanzas", "1"),
    ("Andes Operaciones SpA", "Encargado atencion clientes", "Atencion al Cliente", "2"),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    return table


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for style_name, size, color in [
    ("Heading 1", 16, "2E74B5"),
    ("Heading 2", 13, "2E74B5"),
    ("Heading 3", 12, "1F4D78"),
]:
    style = styles[style_name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(14 if style_name == "Heading 2" else 18)
    style.paragraph_format.space_after = Pt(7 if style_name == "Heading 2" else 10)
    style.paragraph_format.line_spacing = 1.25

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(10)
run = title.add_run("Accesos demo - Perfil Primero")
run.font.name = "Calibri"
run.font.size = Pt(20)
run.font.bold = True
run.font.color.rgb = RGBColor.from_string("07324A")

p = doc.add_paragraph()
p.add_run("Ambiente: ").bold = True
p.add_run("https://perfil-primero.web.app. ")
p.add_run("Valores de prueba: ").bold = True
p.add_run("US$1 para activacion de postulante y US$1 para cierre de empresa.")

doc.add_heading("Credenciales", level=1)
add_table(doc, ["Rol", "Nombre", "Usuario/email", "Password", "Ruta"], credentials)

doc.add_heading("Postulantes demo", level=1)
add_table(doc, ["Nombre", "Perfil", "Habilidades", "Renta"], postulants)

doc.add_heading("Empresas demo", level=1)
add_table(doc, ["Empresa", "RUT demo", "Rubro", "Estado"], [
    ("Norte Digital SpA", "76.345.210-8", "Marketing y Publicidad", "Verificada"),
    ("Andes Operaciones SpA", "77.812.904-1", "Abastecimiento y Logistica", "Verificada"),
])

doc.add_heading("Ofertas demo", level=1)
add_table(doc, ["Empresa", "Cargo", "Area", "Vacantes"], offers)

doc.add_heading("Notas operativas", level=1)
for item in [
    "La consola temporal Admin / 1234 permite entrar visualmente; las acciones reales de verificacion siguen protegidas por Firebase Auth y rol admin.",
    "Las entrevistas deben programarse con al menos un dia de anticipacion y generan enlace de Google Calendar.",
    "Mercado Pago queda configurado para monto de prueba US$1; revisar soporte de moneda USD segun cuenta pais antes de produccion.",
    "Los asientos contables se registran como pendientes de revision, con bruto, IVA 19%, comision estimada Mercado Pago y cuentas sugeridas."
]:
    para = doc.add_paragraph(style=None)
    para.style = styles["Normal"]
    para.paragraph_format.left_indent = Inches(0.375)
    para.paragraph_format.first_line_indent = Inches(-0.188)
    para.add_run("- ")
    para.add_run(item)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer.add_run("Perfil Primero - documento demo")

doc.save(OUT)
print(OUT.resolve())
